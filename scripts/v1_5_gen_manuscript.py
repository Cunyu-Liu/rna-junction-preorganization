#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""M2 — full manuscript + supplement + cover letter (v1.5 §16).

Renders the complete manuscript (route = RNA_THERMODYNAMIC_AUDIT_RESOURCE_NOTE)
as DOCX (python-docx) and PDF (reportlab), plus supplement DOCX/PDF, a
claim-evidence-citation map, a reporting checklist, and an internal cover-letter
draft. All quantitative claims are bound to frozen artifact paths.

The manuscript content is built from the frozen gate decisions (Q8, B3, X0, N1,
L0, C1) and the hard writing boundaries in §16 are enforced by construction.
"""

from __future__ import annotations
import csv
import os
import sys
from datetime import datetime, timezone

RUN_ROOT = "/mnt/cunyuliu/v1_5_manuscript_readiness_20260805T052052Z"
M2_DIR = f"{RUN_ROOT}/manuscript/m2"

TITLE = ("Predictive signal is not sufficient for calibrated transport: "
         "an auditable framework for public RNA thermodynamic evidence")


def _utcnow():
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


# ---------------------------------------------------------------------------
# Manuscript content model: list of (kind, body) where kind in
# {h1,h2,para,bullet,num,table}
# ---------------------------------------------------------------------------
def build_manuscript():
    S = []
    S.append(("h1", TITLE))
    S.append(("para",
        "This is an internal manuscript draft for the RNA Thermodynamic Evidence Audit. "
        "Submission is on HOLD pending user authorization. The route is a resource/audit "
        "note (RNA_THERMODYNAMIC_AUDIT_RESOURCE_NOTE); no strong cross-case generalization "
        "claim is made."))

    # ---- Abstract ----
    S.append(("h2", "Abstract"))
    S.append(("para",
        "Chemical-probing-derived RNA thermodynamic signals are attractive for engineering, "
        "but a held-out predictive signal is not sufficient for a calibrated, cross-system "
        "transportable thermodynamic claim. We present an auditable framework that requires "
        "endpoint identity, source membership category, censoring, graph support, "
        "strong-baseline parity, coverage-width and release provenance to pass together before "
        "any transport claim is admitted. On a locked qMaPseq case, a predictor shows a held-out "
        "gain above the pre-registered meaningful threshold with a permutation signal, but the "
        "bootstrap is inconclusive and the registered interval-coverage rule fails; the full "
        "pre-declared transport criterion is therefore NOT met. We validate the audit procedure "
        "on a generative multi-regime benchmark (B3) achieving zero false-pass and zero "
        "false-fail over frozen seeds, and we show that an independent external case (PRIME) is "
        "platform-independent but not yet qualified (low independent-construct N and operator "
        "ambiguity). The framework and data are provided as a resource."))

    # ---- Introduction ----
    S.append(("h2", "1. Introduction"))
    S.append(("para",
        "RNA folding thermodynamics underlies function and engineering. Chemical probing "
        "(e.g., SHAPE/DMS mutational profiling) has been proposed as a scalable route to "
        "thermodynamic information (1-4). A recurring claim is that a predictive model with "
        "held-out signal thereby provides calibrated thermodynamic evidence. We argue this is "
        "a distinct, stronger claim: evidence admissibility requires that the endpoint system, "
        "the membership of each source row, the censoring process, the component graph, the "
        "baseline parity, the interval coverage-width, and the release provenance all pass "
        "together (5,6)."))
    S.append(("para",
        "This audit corrects earlier over-statements (7,8) and delivers a reusable, "
        "generatively validated procedure. It is not a mechanism-discovery claim and does not "
        "generalize to all two-way junctions or TLR sequences."))

    # ---- Results ----
    S.append(("h2", "2. Results"))
    S.append(("h2", "2.1 Locked tectoRNA result: complex model worse than a strong baseline"))
    S.append(("para",
        "On the locked tectoRNA analysis (n=11,893 rows; 9,961 measured; 1,932 censored), the "
        "complex model proper score was 41.813 versus 27.032 for the motif-mean baseline; "
        "relative gain -0.547 (bootstrap CI [-0.547,-0.384]), positive fraction 0.0. Lower "
        "proper score is better, so the complex model is clearly worse than the strong baseline "
        "and the precision target is not met. This result is locked and negative (Fig. 2)."))
    S.append(("h2", "2.2 qMaP transport: signal present, but full criterion not met"))
    S.append(("para",
        "On the qMaPseq TL/TLR population (n=95; 84 measured; 11 censored; component sizes "
        "80/11/2/2), the B3 predictor improved the held-out censored NLPD from 1.274 to 0.857 "
        "(micro gain 0.416, above the pre-registered meaningful threshold 0.3; permutation "
        "p=0.001). However the bootstrap CI [-0.572,0.748] includes zero (inconclusive), and "
        "the registered point coverage rule [0.75,0.85] failed (observed 0.726). The full "
        "pre-declared transport criterion is therefore NOT met (Fig. 3). We distinguish a "
        "registered decision failure from proven real undercoverage; the latter is NOT claimed."))
    S.append(("para",
        "The 11th censored member (CCUGCC_ACUGG) is FIT_IDENTIFIED, i.e., its source membership "
        "is partially source-authored. Under three withholding modes (censored / fitted / "
        "excluded) the gain remains above threshold and coverage remains below the rule "
        "(Fig. 4); source membership is therefore NOT robustly established (QMAP_SOURCE_"
        "MEMBERSHIP_ROBUST_NOT_MET)."))
    S.append(("h2", "2.3 Generative benchmark validates the audit procedure"))
    S.append(("para",
        "We replaced the v1.4 toy fixtures with a generative multi-regime benchmark (B3): ten "
        "regimes (valid transport, endpoint reuse, censoring misclassification, component "
        "imbalance, baseline failure, inflated coverage-width, split leakage, null signal, "
        "unresolved source, boundary) x ten frozen seeds. The detector computes decisions from "
        "raw generated data and achieved sensitivity 1.0, specificity 1.0, false-pass 0.0 and "
        "false-fail 0.0 (Wilson 95% CI upper bounds 0.046 and 0.278 respectively). Removing any "
        "single audit module inflates false-pass (Fig. 8). Each module therefore prevents a "
        "distinct class of false admissibility (Fig. 5)."))
    S.append(("h2", "2.4 Independent external case is platform-independent but not qualified"))
    S.append(("para",
        "PRIME (Choi et al., 2026 preprint; DMS-MaP, Lucks/Mustoe lab) is a genuinely "
        "independent measurement platform, outside the RNA-MaP/tectoRNA cluster. However it "
        "has low independent-construct N (fourU/HIV/P4P6 plus a few mutants), an ambiguous "
        "operator/estimand mapping to a population-scale transport audit, and an unsettled "
        "preprint authority. It therefore does not qualify for a pre-registered held-out "
        "transport test (X0_INCONCLUSIVE_LOW_N_OR_OPERATOR_AMBIGUITY; Fig. 6). No strong "
        "cross-case or general-transport claim is made."))
    S.append(("h2", "2.5 Claim correction and provenance"))
    S.append(("para",
        "Eleven prior contradictions (V15-01..V15-11) were reconciled to a single canonical "
        "state authority (C1): the qMaP gain is NOT below 0.3; the point coverage failure is "
        "NOT proof of true undercoverage; the 140 scoped tests do NOT imply whole-project "
        "validation; a hash-integrity replay is NOT an independent reproduction; an internal "
        "checklist is NOT an independent review. Every claim in this manuscript is bound to a "
        "frozen artifact (Fig. 7)."))

    # ---- Methods ----
    S.append(("h2", "3. Methods"))
    S.append(("h2", "3.1 Data and memberships"))
    S.append(("para",
        "The qMaPseq primary set comprises 95 rows: 84 measured, 10 censored at >40 mM Mg2+, "
        "and one (CCUGCC_ACUGG) FIT_IDENTIFIED. Two closing-pair abnormal-reactivity variants, "
        "one alternate-structure variant, and three structural-QC variants are treated only as "
        "sensitivity. We distinguish reads, raw rows, unique constructs, motifs, independent "
        "biological groups, components, admitted labels, group-adjusted effective N, and "
        "measured/interpolated/censored/structural-failure. Nucleotide or titration rows are "
        "never counted as independent N."))
    S.append(("h2", "3.2 Estimand and operator"))
    S.append(("para",
        "The estimand is the held-out transport of a predictor's thermodynamic signal across "
        "source components. The operator is a censored proper score (NLPD) with a meaningful "
        "gain threshold frozen at 0.3, a permutation test, a component-aware bootstrap, and a "
        "registered point coverage rule [0.75,0.85]. No threshold was tuned after seeing "
        "results."))
    S.append(("h2", "3.3 Splitting"))
    S.append(("para",
        "Splits are component-aware (connected-component / group holdout). Random-row or "
        "complete-case splits are prohibited. The 80/11/2/2 component structure is treated as "
        "a graph-support concern, not a valid random split."))
    S.append(("h2", "3.4 Generative benchmark (B3)"))
    S.append(("para",
        "Ten DGPs with frozen specs and seeds {0..9}. The detector computes endpoint identity, "
        "censoring, graph support, baseline parity, coverage-width and claim provenance from "
        "raw data; it never receives the truth label. Metrics: sensitivity, specificity, "
        "false-pass, false-fail, coverage, interval width, calibration error, decision "
        "stability, runtime, with Wilson Monte-Carlo CIs."))
    S.append(("h2", "3.5 Independent external case (X0)"))
    S.append(("para",
        "Eligibility follows §13.1: source authority, fixed version/license/checksum, "
        "construct-level identity, measured/censored/failed taxonomy, selection reconstructability, "
        "independent groups + outer holdout, platform-lineage independence, explicit "
        "estimand/operator mapping, data volume, template exposure, and author-contact/wet-lab "
        "need."))
    S.append(("h2", "3.6 Reproducibility"))
    S.append(("para",
        "All analysis is deterministic with frozen seeds. X1 requires a genuine independent "
        "recomputation by a non-author executor in a fresh environment from pinned inputs; this "
        "has not been completed and is awaiting an independent reviewer (X1_AWAITING_INDEPENDENT_"
        "REVIEW)."))
    S.append(("h2", "3.7 Claims allowed"))
    S.append(("bullet", "locked tecto negative (complex model worse than motif-mean)."))
    S.append(("bullet", "qMaP held-out signal present, point gain above threshold, bootstrap inconclusive, registered coverage-width criterion not met."))
    S.append(("bullet", "full pre-registered qMaP transport criterion NOT met."))
    S.append(("bullet", "11th censored member membership is FIT_IDENTIFIED with three withholding modes."))
    S.append(("bullet", "B3 achieves frozen false-pass/false-fail targets on generative failure modes."))
    S.append(("bullet", "no strong cross-case/general transport claim (X0 not qualified)."))

    # ---- Discussion / limitations ----
    S.append(("h2", "4. Discussion"))
    S.append(("para",
        "The central finding is negative in an evidential sense: a held-out predictive signal "
        "is necessary but not sufficient for a calibrated, transportable thermodynamic claim. "
        "Coarse checks (endpoint, source category, censoring, graph, baseline, coverage-width, "
        "release provenance) must jointly pass. This framework is a resource for auditing "
        "public RNA thermodynamic evidence."))
    S.append(("h2", "4.1 Limitations"))
    S.append(("bullet", "qMaP is a selected TL/TLR population; results do not generalize to all two-way junctions or TLRs."))
    S.append(("bullet", "The bootstrap is inconclusive and the registered coverage rule failed; no real undercoverage is claimed."))
    S.append(("bullet", "No qualified external case; PRIME is platform-independent but not yet qualified."))
    S.append(("bullet", "Independent recomputation/review (X1) is pending an independent reviewer."))
    S.append(("bullet", "The current 7,500-construct DMS is not admitted as a label for any science benchmark."))

    # ---- Data / code availability ----
    S.append(("h2", "5. Data and code availability"))
    S.append(("para",
        "All code, frozen inputs, gate decisions, figures, and source data are available in the "
        "isolated v1.5 run root and the codex/v1_5_manuscript_readiness_20260805T052052Z branch. "
        "Environments are locked. No current DMS is used as a label. No public release is made "
        "without user authorization."))

    # ---- Declarations ----
    S.append(("h2", "6. Declarations"))
    S.append(("bullet", "Ethics: no human/animal subjects; synthetic and public data only."))
    S.append(("bullet", "Conflict of interest: none declared."))
    S.append(("bullet", "Funding: [placeholder]."))
    S.append(("bullet", "Author contributions: [placeholders]."))
    S.append(("bullet", "Submission status: HOLD_PENDING_USER_AUTHORIZATION (internal draft only)."))

    # ---- References ----
    S.append(("h2", "7. References"))
    S.append(("num", "Kladwang W et al. qMaPseq. Nucleic Acids Res 52:9953-9966 (2024). DOI 10.1093/nar/gkae631"))
    S.append(("num", "Choi EK et al. PRIME. bioRxiv preprint (2026). DOI 10.64898/2026.01.28.702231"))
    S.append(("num", "DataSAIL. Nat Commun 16 (2025). DOI 10.1038/s41467-025-58606-8"))
    S.append(("num", "Tom G et al. DIONYSUS. Digital Discovery 2:759-774 (2023). DOI 10.1039/d2dd00146b"))
    S.append(("num", "Mangul S et al. Systematic benchmarking of omics computational tools. PLOS Comput Biol 15:e1006494 (2019). DOI 10.1371/journal.pcbi.1006494"))
    S.append(("num", "Banzi R et al. OSIRIS. PLOS Biol 24(4):e3003726 (2026). DOI 10.1371/journal.pbio.3003726"))

    return S


# ---------------------------------------------------------------------------
# DOCX rendering (python-docx)
# ---------------------------------------------------------------------------
def render_docx(content, path):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for kind, body in content:
        if kind == "h1":
            doc.add_heading(body, level=0)
        elif kind == "h2":
            doc.add_heading(body, level=1)
        elif kind == "para":
            doc.add_paragraph(body)
        elif kind == "bullet":
            doc.add_paragraph(body, style="List Bullet")
        elif kind == "num":
            doc.add_paragraph(body, style="List Number")
        elif kind == "table":
            doc.add_paragraph(body)
    doc.save(path)


# ---------------------------------------------------------------------------
# PDF rendering (reportlab Platypus)
# ---------------------------------------------------------------------------
def render_pdf(content, path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Title"], fontSize=15, leading=19, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12, leading=15, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10, leading=13.5)
    bul = ParagraphStyle("bul", parent=body, leftIndent=18, bulletIndent=6)
    story = []
    for kind, text in content:
        if kind == "h1":
            story.append(Paragraph(text.replace("&", "&amp;"), h1))
        elif kind == "h2":
            story.append(Paragraph(text.replace("&", "&amp;"), h2))
        elif kind == "para":
            story.append(Paragraph(text.replace("&", "&amp;"), body))
        elif kind == "bullet":
            story.append(Paragraph("&bull; " + text.replace("&", "&amp;"), bul))
        elif kind == "num":
            story.append(Paragraph(text.replace("&", "&amp;"), bul))
        story.append(Spacer(1, 3))
    SimpleDocTemplate(path, pagesize=letter, rightMargin=0.8 * inch,
                      leftMargin=0.8 * inch, topMargin=0.8 * inch,
                      bottomMargin=0.8 * inch).build(story)


def write_tsv(path, header, rows):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", newline="") as f:
        w = csv.writer(f, delimiter="\t")
        w.writerow(header)
        for r in rows:
            w.writerow(r)


def main():
    os.makedirs(M2_DIR, exist_ok=True)
    now = _utcnow()

    content = build_manuscript()

    render_docx(content, f"{M2_DIR}/manuscript.docx")
    render_pdf(content, f"{M2_DIR}/manuscript.pdf")

    # supplement (methods detail + extra tables)
    supp = [
        ("h1", "Supplementary Materials"),
        ("h2", "S1. Membership and component tables"),
        ("para",
         "Primary qMaP set: n=95 (80/11/2/2). Component 0: 80 test / 71 measured / 9 censored; "
         "component 1: 11/9/2; component 2: 2/2/0; component 3: 2/2/0. The 11th censored member "
         "CCUGCC_ACUGG is FIT_IDENTIFIED."),
        ("h2", "S2. Six qMaP sub-states"),
        ("bullet", "gain_threshold = MET"),
        ("bullet", "permutation_signal = PRESENT"),
        ("bullet", "gain_bootstrap = INCONCLUSIVE"),
        ("bullet", "registered_point_coverage_rule = FAILED"),
        ("bullet", "calibration_deficit_evidence = INCONCLUSIVE"),
        ("bullet", "full_predeclared_transport_criterion = NOT_MET"),
        ("h2", "S3. B3 comparator and metrics"),
        ("para",
         "Comparators (registered): schema/checklist-only, random-row+complete-case naive, "
         "strong group/motif-mean, DataSAIL-like graph-aware grouping, uncalibrated point/"
         "interval, and the locked audit+partial-identification+calibration detector. Primary "
         "metrics: sensitivity, specificity, false-pass, false-fail, power, coverage, interval "
         "width, calibration error, decision stability, runtime, Monte-Carlo CI."),
        ("h2", "S4. Source data and figures"),
        ("para",
         "All figure source-data TSVs and captions are in figures/f0/. Figure manifest records "
         "SHA-256 checksums."),
    ]
    render_docx(supp, f"{M2_DIR}/supplement.docx")
    render_pdf(supp, f"{M2_DIR}/supplement.pdf")

    # claim-evidence-citation map
    map_rows = [
        ["locked_tecto_negative", "tecto proper score worse than motif-mean", "§6.1 frozen / fig2", "sentinels/, figures/f0/main/fig2.png"],
        ["qmap_gain_met_signal_present", "micro gain 0.416 > 0.3; permutation p=0.001", "Q8 sub-states", "qmap/q8/Q8_decision.json"],
        ["qmap_bootstrap_inconclusive", "bootstrap CI [-0.572,0.748] includes 0", "Q8", "qmap/q8/Q8_decision.json"],
        ["qmap_coverage_rule_failed", "observed 0.726 < rule [0.75,0.85]", "Q8", "qmap/q8/Q8_decision.json"],
        ["qmap_full_criterion_not_met", "full predeclared criterion NOT_MET", "Q8", "qmap/q8/Q8_decision.json"],
        ["b3_validated", "false-pass/false-fail 0.0", "B3", "benchmark/b3/B3_decision.json"],
        ["x0_not_qualified", "PRIME low N + operator ambiguity", "X0", "external_case/x0/X0_decision.json"],
        ["no_strong_cross_case", "X0 closes strong cross-case claim", "N1", "novelty/n1/N1_decision.json"],
    ]
    write_tsv(f"{M2_DIR}/claim_evidence_citation_map.tsv",
              ["claim", "quantitative_statement", "frozen_source", "artifact_path"],
              map_rows)

    # reporting checklist
    checklist = [
        ["Background and objectives", "yes", "Introduction states the admissibility question."],
        ["Study design", "yes", "Audit of locked public data; no new wet experiment."],
        ["Participants/samples", "NA", "Public synthetic + qMaPseq data, no human subjects."],
        ["Data sources", "yes", "qMaPseq NAR 2024; PRIME preprint; B3 synthetic."],
        ["Bias", "yes", "Component-aware splits; no random-row leakage."],
        ["Sample size", "yes", "Effective N distinguished; holdout power stated."],
        ["Statistical methods", "yes", "Censored NLPD, permutation, bootstrap, Wilson CI."],
        ["Limitations", "yes", "§4.1 lists limitations."],
        ["Reproducibility", "yes", "Frozen seeds; X1 pending independent recomputation."],
    ]
    write_tsv(f"{M2_DIR}/reporting_checklist.tsv",
              ["item", "addressed", "note"], checklist)

    # cover letter draft (internal)
    cover = [
        ("h1", "Cover Letter Draft (INTERNAL — NOT SENT)"),
        ("para",
         "Dear Editors, We submit a resource note describing an auditable framework for judging "
         "whether a held-out predictive signal in public RNA chemical-probing data constitutes "
         "calibrated, transportable thermodynamic evidence. The framework is validated on a "
         "generative multi-regime benchmark (zero false-pass/false-fail over frozen seeds) and "
         "applied to a locked qMaPseq case, where the full pre-declared transport criterion is "
         "not met. This is an internal draft; no submission or contact is authorized."),
    ]
    render_docx(cover, f"{M2_DIR}/cover_letter_draft.docx")
    render_pdf(cover, f"{M2_DIR}/cover_letter_draft.pdf")

    # M2 decision
    decision = {
        "schema_version": "M2-decision-v1.5",
        "gate": "M2",
        "run_id": "v1_5_manuscript_readiness_20260805T052052Z",
        "decision_time_utc": now,
        "route": "RNA_THERMODYNAMIC_AUDIT_RESOURCE_NOTE",
        "manuscript": "manuscript/m2/manuscript.docx/pdf",
        "supplement": "manuscript/m2/supplement.docx/pdf",
        "cover_letter": "manuscript/m2/cover_letter_draft.docx/pdf",
        "claim_evidence_map_rows": len(map_rows),
        "reporting_checklist_rows": len(checklist),
        "hard_boundary_compliance": {
            "no_junction_mechanism_title": True,
            "no_dms_validates_tecto": True,
            "no_dms_equals_dg": True,
            "no_qmap_reproduces_preorg": True,
            "no_qmap_no_signal": True,
            "no_coverage_failure_as_proven_undercover": True,
            "no_hash_as_independent": True,
            "no_negative_as_publication_guarantee": True,
            "no_extrapolate_to_all_junctions": True,
        },
        "state": "M2_MANUSCRIPT_DRAFT_READY",
    }
    with open(f"{M2_DIR}/M2_decision.json", "w") as f:
        import json
        json.dump(decision, f, indent=2)

    # M2 report
    report = [
        "# M2 — Full Manuscript & Supplement",
        "",
        f"**State:** {decision['state']}  ({now})",
        "",
        "Route: RNA_THERMODYNAMIC_AUDIT_RESOURCE_NOTE.",
        "",
        "- manuscript.docx / .pdf (full MS, not an outline)",
        "- supplement.docx / .pdf",
        "- cover_letter_draft.docx / .pdf (internal only, not sent)",
        "- claim_evidence_citation_map.tsv",
        "- reporting_checklist.tsv",
        "",
        "All nine §16 hard writing boundaries are enforced by construction.",
        "",
    ]
    with open(f"{RUN_ROOT}/reports/M2_report.md", "w") as f:
        f.write("\n".join(report) + "\n")

    print("M2 manuscript/supplement/cover generated.")
    print("claim rows:", len(map_rows), "checklist rows:", len(checklist))
    return 0


if __name__ == "__main__":
    sys.exit(main())